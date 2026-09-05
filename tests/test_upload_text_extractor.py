import io

from docx import Document

from backend.services.upload_text_extractor import extract_uploaded_text


def test_extracts_text_and_docx_without_persisting_private_upload():
    assert extract_uploaded_text(
        b"private markdown", filename="note.md", content_type="text/markdown",
    ) == "private markdown"
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Private method", level=1)
    document.add_paragraph("Verified outcome")
    document.save(buffer)
    extracted = extract_uploaded_text(
        buffer.getvalue(), filename="method.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Private method" in extracted and "Verified outcome" in extracted
